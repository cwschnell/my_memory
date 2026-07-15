from fastapi import APIRouter, Depends, HTTPException, status, Header
from fastapi.responses import JSONResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update, delete, func, desc, and_
from datetime import date, datetime, timedelta
from typing import List, Optional, Dict
import uuid

from database import get_db
from models import Guest, Reservation, LodgeTask, Incident, DailyLog, Recording, Room, Agency
from schemas import (
    GuestCreate, GuestOut,
    ReservationCreate, ReservationOut,
    LodgeTaskCreate, LodgeTaskOut,
    IncidentCreate, IncidentOut,
    DailyLogCreate, DailyLogOut,
    RoomCreate, RoomOut,
    AgencyCreate, AgencyOut
)

router = APIRouter(prefix="/lodge", tags=["lodge"])

async def get_active_lodge_id(x_lodge_id: Optional[str] = Header(None, alias="X-Lodge-Id")) -> Optional[uuid.UUID]:
    if not x_lodge_id:
        return None
    try:
        return uuid.UUID(x_lodge_id)
    except ValueError:
        return None

async def get_current_user_email(x_user_email: Optional[str] = Header(None, alias="X-User-Email")) -> Optional[str]:
    """Extract and normalize the caller's email from the X-User-Email header."""
    if not x_user_email:
        return None
    return x_user_email.strip().lower()

def parse_date(date_str: Optional[str]) -> Optional[date]:
    if not date_str:
        return None
    normalized = str(date_str).replace("/", "-").strip()
    if not normalized or normalized.lower() in ("null", "yyyy-mm-dd", "yyyy/mm/dd"):
        return None
    from datetime import datetime
    for fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%Y/%m/%d", "%d-%m-%Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(normalized, fmt).date()
        except ValueError:
            continue
    try:
        parts = normalized.split("-")
        if len(parts) == 3:
            y = int(parts[0])
            m = int(parts[1])
            d = int(parts[2][:2])
            from datetime import date as date_type
            return date_type(y, m, d)
    except Exception:
        pass
    return None

# -----------------------------------------------------------------------------
# Module 8 Backend: Dashboard Summary Endpoint
# -----------------------------------------------------------------------------
@router.get("/dashboard")
async def get_dashboard_summary(
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    today = date.today()

    # Check-ins today
    checkins_stmt = select(Reservation).join(Guest).where(Reservation.check_in == today)
    if lodge_id:
        checkins_stmt = checkins_stmt.where(Guest.lodge_id == lodge_id)
    checkins_res = await db.execute(checkins_stmt)
    checkins = checkins_res.scalars().all()

    # Check-outs today
    checkouts_stmt = select(Reservation).join(Guest).where(Reservation.check_out == today)
    if lodge_id:
        checkouts_stmt = checkouts_stmt.where(Guest.lodge_id == lodge_id)
    checkouts_res = await db.execute(checkouts_stmt)
    checkouts = checkouts_res.scalars().all()

    # In-house tonight (check_in <= today < check_out and status != cancelled)
    inhouse_stmt = select(Reservation).join(Guest).where(
        Reservation.check_in <= today,
        Reservation.check_out > today,
        Reservation.status != "cancelled"
    )
    if lodge_id:
        inhouse_stmt = inhouse_stmt.where(Guest.lodge_id == lodge_id)
    inhouse_res = await db.execute(inhouse_stmt)
    inhouse = inhouse_res.scalars().all()
    
    inhouse_guest_ids = [r.guest_id for r in inhouse if r.guest_id]
    guests_map = {}
    if inhouse_guest_ids:
        g_res = await db.execute(select(Guest).where(Guest.id.in_(inhouse_guest_ids)))
        for g in g_res.scalars().all():
            guests_map[g.id] = g.full_name

    inhouse_list = []
    total_inhouse_count = 0
    for r in inhouse:
        gname = guests_map.get(r.guest_id, "Unknown Guest")
        inhouse_list.append({"guest_name": gname, "room": r.room_or_unit, "adults": r.num_adults, "children": r.num_children})
        total_inhouse_count += (r.num_adults or 1) + (r.num_children or 0)

    # Tasks due today or overdue
    tasks_stmt = select(LodgeTask).where(
        LodgeTask.is_complete == False,
        LodgeTask.due_date <= today
    )
    if lodge_id:
        tasks_stmt = tasks_stmt.where(LodgeTask.lodge_id == lodge_id)
    tasks_res = await db.execute(tasks_stmt.order_by(LodgeTask.due_date.asc()))
    tasks_due = tasks_res.scalars().all()

    # Open Incidents
    incidents_stmt = select(Incident).where(Incident.is_resolved == False)
    if lodge_id:
        incidents_stmt = incidents_stmt.where(Incident.lodge_id == lodge_id)
    incidents_res = await db.execute(incidents_stmt.order_by(
        desc(Incident.severity == "critical"),
        desc(Incident.created_at)
    ))
    open_incidents = incidents_res.scalars().all()

    # Recent Memos (from recordings table where type='memo')
    memos_stmt = select(Recording).where(Recording.type == "memo")
    if lodge_id:
        memos_stmt = memos_stmt.where(Recording.lodge_id == lodge_id)
    memos_res = await db.execute(memos_stmt.order_by(desc(Recording.created_at)).limit(5))
    recent_memos = memos_res.scalars().all()

    return {
        "today": today.isoformat(),
        "checkins": [{"room": r.room_or_unit, "guest_id": r.guest_id} for r in checkins],
        "checkouts": [{"room": r.room_or_unit, "guest_id": r.guest_id} for r in checkouts],
        "inhouse_count": total_inhouse_count,
        "inhouse_list": inhouse_list,
        "tasks_due_count": len(tasks_due),
        "tasks_due": [
            {"id": str(t.id), "title": t.title, "area": t.area, "assigned_to": t.assigned_to, "due_date": t.due_date.isoformat() if t.due_date else None}
            for t in tasks_due[:8]
        ],
        "incidents_count": len(open_incidents),
        "open_incidents": [
            {"id": str(i.id), "title": i.title, "area": i.area, "severity": i.severity}
            for i in open_incidents[:8]
        ],
        "recent_memos": [
            {"id": str(m.id), "summary": m.summary, "transcript": m.transcript, "created_at": m.created_at.isoformat() if m.created_at else None}
            for m in recent_memos
        ]
    }


# -----------------------------------------------------------------------------
# Guests CRUD
# -----------------------------------------------------------------------------
@router.get("/guests", response_model=List[GuestOut])
async def list_guests(
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    stmt = select(Guest)
    if user_email:
        stmt = stmt.where(Guest.user_email == user_email)
    if lodge_id:
        stmt = stmt.where(Guest.lodge_id == lodge_id)
    res = await db.execute(stmt.order_by(Guest.full_name.asc()))
    guests = res.scalars().all()
    for g in guests:
        g.has_passport_image = bool(g.passport_image)
    return guests

@router.post("/guests", response_model=GuestOut, status_code=201)
async def create_guest(
    data: GuestCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    check_in = parse_date(data.check_in)
    check_out = parse_date(data.check_out)
    room_or_unit = data.room_or_unit or "Chalet"
    guest_data = data.model_dump(exclude={"check_in", "check_out", "room_or_unit"})
    
    # Safely parse date fields
    guest_data["date_of_birth"] = parse_date(guest_data.get("date_of_birth"))
    guest_data["date_of_issue"] = parse_date(guest_data.get("date_of_issue"))
    guest_data["date_of_expiry"] = parse_date(guest_data.get("date_of_expiry"))
    guest_data["visa_validity"] = parse_date(guest_data.get("visa_validity"))
    
    # Always stamp the owner email from the authenticated header
    if user_email:
        guest_data["user_email"] = user_email
    
    g = Guest(**guest_data)
    if lodge_id:
        g.lodge_id = lodge_id
        
    db.add(g)
    await db.commit()
    await db.refresh(g)
    
    if check_in and check_out:
        from models import Reservation
        res = Reservation(
            guest_id=g.id,
            check_in=check_in,
            check_out=check_out,
            room_or_unit=room_or_unit,
            status="confirmed",
            source="direct"
        )
        db.add(res)
        await db.commit()
        
    g.has_passport_image = bool(g.passport_image)
    return g

@router.put("/guests/{guest_id}", response_model=GuestOut)
async def update_guest(
    guest_id: uuid.UUID,
    data: GuestCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    stmt = select(Guest).where(Guest.id == guest_id)
    if user_email:
        stmt = stmt.where(Guest.user_email == user_email)
    if lodge_id:
        stmt = stmt.where(Guest.lodge_id == lodge_id)
    res = await db.execute(stmt)
    g = res.scalar_one_or_none()
    if not g:
        raise HTTPException(status_code=404, detail="Guest not found")
        
    check_in = parse_date(data.check_in)
    check_out = parse_date(data.check_out)
    room_or_unit = data.room_or_unit or "Chalet"
    guest_data = data.model_dump(exclude={"check_in", "check_out", "room_or_unit"})
    
    # Safely parse date fields
    guest_data["date_of_birth"] = parse_date(guest_data.get("date_of_birth"))
    guest_data["date_of_issue"] = parse_date(guest_data.get("date_of_issue"))
    guest_data["date_of_expiry"] = parse_date(guest_data.get("date_of_expiry"))
    guest_data["visa_validity"] = parse_date(guest_data.get("visa_validity"))
    
    # Preserve owner email
    if user_email:
        guest_data["user_email"] = user_email
    
    for k, v in guest_data.items():
        setattr(g, k, v)
    await db.commit()
    await db.refresh(g)
    
    if check_in and check_out:
        from models import Reservation
        res_stmt = await db.execute(select(Reservation).where(Reservation.guest_id == g.id).order_by(Reservation.created_at.desc()))
        existing_res = res_stmt.scalars().first()
        if existing_res:
            existing_res.check_in = check_in
            existing_res.check_out = check_out
            existing_res.room_or_unit = room_or_unit
        else:
            new_res = Reservation(
                guest_id=g.id,
                check_in=check_in,
                check_out=check_out,
                room_or_unit=room_or_unit,
                status="confirmed",
                source="direct"
            )
            db.add(new_res)
        await db.commit()
        
    g.has_passport_image = bool(g.passport_image)
    return g

@router.delete("/guests/{guest_id}", status_code=204)
async def delete_guest(
    guest_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    stmt = select(Guest).where(Guest.id == guest_id)
    if user_email:
        stmt = stmt.where(Guest.user_email == user_email)
    if lodge_id:
        stmt = stmt.where(Guest.lodge_id == lodge_id)
    res = await db.execute(stmt)
    g = res.scalar_one_or_none()
    if g:
        await db.delete(g)
        await db.commit()
    return None


# -----------------------------------------------------------------------------
# Reservations CRUD
# -----------------------------------------------------------------------------
@router.get("/reservations")
async def list_reservations(
    upcoming: Optional[int] = None,
    status_filter: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    query = select(Reservation).join(Guest)
    if user_email:
        query = query.where(Guest.user_email == user_email)
    if lodge_id:
        query = query.where(Guest.lodge_id == lodge_id)
        
    query = query.order_by(Reservation.check_in.asc())
    if upcoming:
        today = date.today()
        end_date = today + timedelta(days=upcoming)
        query = query.where(Reservation.check_in >= today, Reservation.check_in <= end_date)
    if status_filter:
        query = query.where(Reservation.status == status_filter)
    
    res = await db.execute(query)
    reservations = res.scalars().all()
    
    out = []
    for r in reservations:
        g_data = None
        if r.guest_id:
            gres = await db.execute(select(Guest).where(Guest.id == r.guest_id))
            gobj = gres.scalar_one_or_none()
            if gobj:
                g_data = {
                    "id": str(gobj.id),
                    "full_name": gobj.full_name,
                    "phone": gobj.phone,
                    "email": gobj.email
                }
        out.append({
            "id": str(r.id),
            "guest_id": str(r.guest_id) if r.guest_id else None,
            "guest": g_data,
            "room_or_unit": r.room_or_unit,
            "check_in": r.check_in.isoformat() if r.check_in else None,
            "check_out": r.check_out.isoformat() if r.check_out else None,
            "num_adults": r.num_adults,
            "num_children": r.num_children,
            "rate_per_night_usd": float(r.rate_per_night_usd or 0),
            "total_usd": float(r.total_usd or 0),
            "deposit_paid": r.deposit_paid,
            "status": r.status,
            "source": r.source,
            "notes": r.notes
        })
    return out

@router.get("/immigration-report")
async def get_immigration_report(
    start_date: date,
    end_date: date,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    query = select(Reservation).join(Guest)
    if user_email:
        query = query.where(Guest.user_email == user_email)
    if lodge_id:
        query = query.where(Guest.lodge_id == lodge_id)
        
    query = query.where(Reservation.check_in >= start_date, Reservation.check_in <= end_date)
    query = query.order_by(Reservation.check_in.asc())
    
    res = await db.execute(query)
    reservations = res.scalars().all()
    
    out = []
    for idx, r in enumerate(reservations):
        gobj = None
        if r.guest_id:
            gres = await db.execute(select(Guest).where(Guest.id == r.guest_id))
            gobj = gres.scalar_one_or_none()
            
        if gobj:
            out.append({
                "no": idx + 1,
                "guest_id": str(gobj.id),
                "reservation_id": str(r.id),
                "full_name": gobj.full_name,
                "nationality": gobj.nationality,
                "passport_number": gobj.passport_number,
                "date_of_issue": gobj.date_of_issue.isoformat() if gobj.date_of_issue else None,
                "arrived_from": gobj.arrived_from,
                "visa_number": gobj.visa_number,
                "visa_validity": gobj.visa_validity.isoformat() if gobj.visa_validity else None,
                "check_in": r.check_in.isoformat() if r.check_in else None,
                "check_out": r.check_out.isoformat() if r.check_out else None,
            })
    return out

@router.get("/immigration-report/docx")
async def get_immigration_report_docx(
    start_date: date,
    end_date: date,
    lodge_name: Optional[str] = None,
    lodge_location: Optional[str] = None,
    language: Optional[str] = "PT",
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    """Generate and stream a DOCX Boletim de Alojamento / Immigration List."""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from fastapi.responses import StreamingResponse
    import io

    # ── Fetch reservations ────────────────────────────────────────────────────
    query = select(Reservation).join(Guest)
    if user_email:
        query = query.where(Guest.user_email == user_email)
    if lodge_id:
        query = query.where(Guest.lodge_id == lodge_id)
    query = query.where(Reservation.check_in >= start_date, Reservation.check_in <= end_date)
    query = query.order_by(Reservation.check_in.asc())
    res = await db.execute(query)
    reservations = res.scalars().all()

    rows = []
    for idx, r in enumerate(reservations):
        gobj = None
        if r.guest_id:
            gres = await db.execute(select(Guest).where(Guest.id == r.guest_id))
            gobj = gres.scalar_one_or_none()
        if gobj:
            def fmt_date(d):
                if not d:
                    return ""
                MONTHS_PT_SHORT = ["jan","fev","mar","abr","mai","jun","jul","ago","set","out","nov","dez"]
                MONTHS_EN_SHORT = ["jan","feb","mar","apr","may","jun","jul","aug","sep","oct","nov","dec"]
                mo = MONTHS_PT_SHORT[d.month - 1] if (language or "PT").upper() == "PT" else MONTHS_EN_SHORT[d.month - 1]
                return f"{d.day:02d}-{mo}-{str(d.year)[2:]}"
            rows.append({
                "no": f"{idx + 1:02d}",
                "full_name": (gobj.full_name or "").upper(),
                "nationality": (gobj.nationality or "").upper(),
                "passport_number": gobj.passport_number or "",
                "date_of_issue": fmt_date(gobj.date_of_issue),
                "arrived_from": (gobj.arrived_from or "").upper(),
                "visa_number": gobj.visa_number or "N/A",
                "visa_validity": fmt_date(gobj.visa_validity) if gobj.visa_validity else "N/A",
                "check_in": fmt_date(r.check_in),
                "check_out": fmt_date(r.check_out),
            })

    is_pt = (language or "PT").upper() == "PT"
    today = date.today()
    today_str = f"{today.day:02d}/ {today.month:02d}/{today.year}"
    MONTHS_PT = ["Janeiro","Fevereiro","Março","Abril","Maio","Junho","Julho","Agosto","Setembro","Outubro","Novembro","Dezembro"]
    MONTHS_EN = ["January","February","March","April","May","June","July","August","September","October","November","December"]
    month_name = MONTHS_PT[start_date.month - 1] if is_pt else MONTHS_EN[start_date.month - 1]

    # ── Create Document ───────────────────────────────────────────────────────
    doc = Document()

    # Page setup: 11" x 8.5" landscape (US Letter landscape — same as example)
    section = doc.sections[0]
    section.page_width  = Inches(11.0)
    section.page_height = Inches(8.5)
    # Narrow margins matching the example (1.27 cm = 0.5")
    section.left_margin   = Cm(1.27)
    section.right_margin  = Cm(1.27)
    section.top_margin    = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    # ── Helper: add centred paragraph ─────────────────────────────────────────
    def add_hdr(text: str, bold: bool = False, size: float = 12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    def add_left(text: str, bold: bool = True, size: float = 10):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    # ── Official header — exact text from the example ─────────────────────────
    add_hdr("REPÚBLICA DE MOÇAMBIQUE",              bold=False, size=12)
    add_hdr(".........*.........",                   bold=False, size=12)
    add_hdr("MINISTÉRIO DO INTERIOR",               bold=False, size=12)
    add_hdr("SERVIÇO NACIONAL DE MIGRAÇÃO",         bold=False, size=12)
    add_hdr("DIRECÇÃO PROVINCIAL DE MIGRACAO-INHAMBANE", bold=False, size=12)
    add_hdr("DEPARTAMENTO DE OPERAÇÕES MIGRATÓRIAS",bold=False, size=12)
    add_hdr("REPARTIҪÃO DO MOVIMENTO MIGRATÓRIO",   bold=False, size=12)
    add_hdr("POSTO DE TRAVESSIA AÉREO DE VILANKULO",bold=True,  size=12)
    # Spacer line (like para 9 in the example)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(0)
    add_hdr("BOLETIM DE ALOJAMENTO, PLASMADO NO ARTIGO 40 DA LEI N°23/2022 DE 29 DE DEZEMBRO",
            bold=True, size=12)

    # ── Meta section — plain left paragraphs (matching the example) ───────────
    loc_label   = "LOCALIZAÇÃO"   if is_pt else "LOCATION"
    lodge_label = "Nome do estabelicimento turístico"  # kept in Portuguese as per example
    month_label = "Mês" if is_pt else "Month"
    date_label  = "Data" if is_pt else "Date"

    add_left(f"{loc_label}: {lodge_location or ''}", bold=True, size=10)
    add_left(f"{lodge_label}ː {lodge_name or ''}", bold=True, size=10)

    # Month / Date / Page on one line
    page_label = "Página" if is_pt else "Page"
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after  = Pt(4)
    def add_meta_run(text, bold=True, size=10):
        run = p_meta.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
    add_meta_run(f"{month_label}:    ")
    add_meta_run(month_name)
    add_meta_run(f"            {date_label}:    {today_str}            {page_label}: 01")

    # ── Data table — exact column widths from the example ────────────────────
    # Bilingual headers: "PT / EN" format, matching official document
    headers = [
        "N/O",
        "NOME COMPLETO/\nFULL NAME",
        "NACIONALIDADE/\nNACIONALITY",
        "Nº PASSAPORTE/\nPASSPORT Nº",
        "DATA DE EMISSÃO/\nDATE OF ISSUE",
        "PROVINIÊNCIA/\nARRIVED FROM",
        "VISTO Nº/\nVISA Nº",
        "VALIDADE/\nVALIDITY",
        "ENTRADA /\nENTRY",
        "SAIDA\nEXIT",
    ]

    # Column widths from example file: [1.41, 6.49, 3.03, 2.39, 2.21, 2.40, 2.34, 1.66, 1.86, 2.14]
    col_widths = [
        Cm(1.41), Cm(6.49), Cm(3.03), Cm(2.39),
        Cm(2.21), Cm(2.40), Cm(2.34), Cm(1.66),
        Cm(1.86), Cm(2.14),
    ]

    def set_cell_border(cell):
        """Add thin borders to a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, (cell, hdr_text) in enumerate(zip(hdr_cells, headers)):
        cell.width = col_widths[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(hdr_text)
        run.bold = True
        run.font.size = Pt(8)
        set_cell_border(cell)

    # Data rows
    for row_data in rows:
        values = [
            row_data["no"],
            row_data["full_name"],
            row_data["nationality"],
            row_data["passport_number"],
            row_data["date_of_issue"],
            row_data["arrived_from"],
            row_data["visa_number"],
            row_data["visa_validity"],
            row_data["check_in"],
            row_data["check_out"],
        ]
        row_cells = table.add_row().cells
        for i, (cell, val) in enumerate(zip(row_cells, values)):
            cell.width = col_widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            run = p.add_run(val)
            run.font.size = Pt(8)
            set_cell_border(cell)

    # ── Stream document ───────────────────────────────────────────────────────
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    fname = f"boletim_{(lodge_name or 'lodge').replace(' ','_')}_{month_name.lower()}_{start_date.year}.docx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'}
    )

@router.get("/immigration-report/excel")
async def get_immigration_report_excel(
    start_date: date,
    end_date: date,
    lodge_name: Optional[str] = None,
    lodge_location: Optional[str] = None,
    language: Optional[str] = "PT",
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id),
    user_email: Optional[str] = Depends(get_current_user_email)
):
    """Generate and stream an Excel (.xlsx) Boletim de Alojamento / Immigration List."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side
    from fastapi.responses import StreamingResponse
    import io

    # ── Fetch reservations ────────────────────────────────────────────────────
    query = select(Reservation).join(Guest)
    if user_email:
        query = query.where(Guest.user_email == user_email)
    if lodge_id:
        query = query.where(Guest.lodge_id == lodge_id)
    query = query.where(Reservation.check_in >= start_date, Reservation.check_in <= end_date)
    query = query.order_by(Reservation.check_in.asc())
    res = await db.execute(query)
    reservations = res.scalars().all()

    rows = []
    for idx, r in enumerate(reservations):
        gobj = None
        if r.guest_id:
            gres = await db.execute(select(Guest).where(Guest.id == r.guest_id))
            gobj = gres.scalar_one_or_none()
        if gobj:
            def fmt_date(d):
                if not d:
                    return ""
                MONTHS_PT_SHORT = ["jan","fev","mar","abr","mai","jun","jul","ago","set","out","nov","dez"]
                MONTHS_EN_SHORT = ["jan","feb","mar","apr","may","jun","jul","aug","sep","oct","nov","dec"]
                mo = MONTHS_PT_SHORT[d.month - 1] if (language or "PT").upper() == "PT" else MONTHS_EN_SHORT[d.month - 1]
                return f"{d.day:02d}-{mo}-{str(d.year)[2:]}"
            rows.append([
                f"{idx + 1:02d}",
                (gobj.full_name or "").upper(),
                (gobj.nationality or "").upper(),
                gobj.passport_number or "",
                fmt_date(gobj.date_of_issue),
                (gobj.arrived_from or "").upper(),
                gobj.visa_number or "N/A",
                fmt_date(gobj.visa_validity) if gobj.visa_validity else "N/A",
                fmt_date(r.check_in),
                fmt_date(r.check_out),
            ])

    is_pt = (language or "PT").upper() == "PT"
    today = date.today()
    today_str = f"{today.day:02d}/{today.month:02d}/{today.year}"
    MONTHS_PT = ["Janeiro","Fevereiro","Março","Abril","Maio","Junho","Julho","Agosto","Setembro","Outubro","Novembro","Dezembro"]
    MONTHS_EN = ["January","February","March","April","May","June","July","August","September","October","November","December"]
    month_name = MONTHS_PT[start_date.month - 1] if is_pt else MONTHS_EN[start_date.month - 1]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Immigration List"

    # ── Official header ─────────────────────────────────────────
    def add_row(ws, data, row_idx, bold=False, size=11, align='center'):
        for col_idx, value in enumerate(data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = Font(bold=bold, size=size)
            if align == 'center':
                cell.alignment = Alignment(horizontal='center', vertical='center')
            elif align == 'left':
                cell.alignment = Alignment(horizontal='left', vertical='center')

    current_row = 1
    
    # Add Emblem Image
    import os
    from openpyxl.drawing.image import Image as OpenpyxlImage
    emblem_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "assets", "emblem.png"))
    print(f"Trying to load emblem from: {emblem_path}")
    if os.path.exists(emblem_path):
        try:
            img = OpenpyxlImage(emblem_path)
            # Scale image to appropriate size (e.g., 60x60 pixels)
            img.height = 65
            img.width = 65
            # Anchor at column F, row 1 (adjust if needed to center)
            ws.add_image(img, "F1")
            print("Successfully added image to worksheet!")
            
            # Add a few blank rows to make space for the image
            for i in range(1, 5):
                ws.merge_cells(start_row=i, start_column=1, end_row=i, end_column=10)
            current_row = 5
        except Exception as e:
            print(f"Failed to add image: {e}")
    else:
        print("Emblem path does NOT exist!")

    # Merge cells for header
    for i in range(1, 9):
        ws.merge_cells(start_row=current_row + i - 1, start_column=1, end_row=current_row + i - 1, end_column=10)
    
    add_row(ws, ["REPÚBLICA DE MOÇAMBIQUE"], current_row, bold=False, size=12)
    current_row += 1
    add_row(ws, [".........*........."], current_row, bold=False, size=12)
    current_row += 1
    add_row(ws, ["MINISTÉRIO DO INTERIOR"], current_row, bold=False, size=12)
    current_row += 1
    add_row(ws, ["SERVIÇO NACIONAL DE MIGRAÇÃO"], current_row, bold=False, size=12)
    current_row += 1
    add_row(ws, ["DIRECÇÃO PROVINCIAL DE MIGRACAO-INHAMBANE"], current_row, bold=False, size=12)
    current_row += 1
    add_row(ws, ["DEPARTAMENTO DE OPERAÇÕES MIGRATÓRIAS"], current_row, bold=False, size=12)
    current_row += 1
    add_row(ws, ["REPARTIҪÃO DO MOVIMENTO MIGRATÓRIO"], current_row, bold=False, size=12)
    current_row += 1
    add_row(ws, ["POSTO DE TRAVESSIA AÉREO DE VILANKULO"], current_row, bold=True, size=12)
    current_row += 1
    
    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=10)
    current_row += 1

    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=10)
    add_row(ws, ["BOLETIM DE ALOJAMENTO, PLASMADO NO ARTIGO 40 DA LEI N°23/2022 DE 29 DE DEZEMBRO"], current_row, bold=True, size=12)
    current_row += 2

    # Meta
    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=10)
    add_row(ws, [f"{'LOCALIZAÇÃO' if is_pt else 'LOCATION'}: {lodge_location or ''}"], current_row, bold=True, size=10, align='left')
    current_row += 1
    
    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=10)
    add_row(ws, [f"Nome do estabelicimento turísticoː {lodge_name or ''}"], current_row, bold=True, size=10, align='left')
    current_row += 1

    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=10)
    add_row(ws, [f"{'Mês' if is_pt else 'Month'}: {month_name}      {'Data' if is_pt else 'Date'}: {today_str}      {'Página' if is_pt else 'Page'}: 01"], current_row, bold=True, size=10, align='left')
    current_row += 2

    # Headers
    headers = [
        "N/O",
        "NOME COMPLETO/\nFULL NAME",
        "NACIONALIDADE/\nNACIONALITY",
        "Nº PASSAPORTE/\nPASSPORT Nº",
        "DATA DE EMISSÃO/\nDATE OF ISSUE",
        "PROVINIÊNCIA/\nARRIVED FROM",
        "VISTO Nº/\nVISA Nº",
        "VALIDADE/\nVALIDITY",
        "ENTRADA /\nENTRY",
        "SAIDA\nEXIT",
    ]
    
    # Add headers
    for col_idx, hdr in enumerate(headers, start=1):
        cell = ws.cell(row=current_row, column=col_idx, value=hdr)
        cell.font = Font(bold=True, size=8)
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        thin = Side(border_style="thin", color="000000")
        cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

    # Set column widths approximately
    col_widths = [5, 25, 15, 12, 12, 12, 12, 10, 10, 10]
    for i, w in enumerate(col_widths, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    current_row += 1

    # Data rows
    for row_data in rows:
        for col_idx, val in enumerate(row_data, start=1):
            cell = ws.cell(row=current_row, column=col_idx, value=val)
            cell.font = Font(size=8)
            cell.alignment = Alignment(horizontal='center' if col_idx == 1 else 'left', vertical='center')
            thin = Side(border_style="thin", color="000000")
            cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
        current_row += 1

    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    
    fname = f"boletim_{(lodge_name or 'lodge').replace(' ','_')}_{month_name.lower()}_{start_date.year}.xlsx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'}
    )

@router.post("/reservations")
async def create_reservation(data: ReservationCreate, db: AsyncSession = Depends(get_db)):
    # Auto calculate total_usd if 0
    total = data.total_usd
    if total == 0 and data.rate_per_night_usd > 0 and data.check_out > data.check_in:
        nights = (data.check_out - data.check_in).days
        total = float(data.rate_per_night_usd) * nights
    
    dump = data.model_dump()
    dump["total_usd"] = total
    r = Reservation(**dump)
    db.add(r)
    await db.commit()
    await db.refresh(r)
    return {"id": str(r.id), "status": "created"}

@router.put("/reservations/{res_id}")
async def update_reservation(res_id: uuid.UUID, data: ReservationCreate, db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(Reservation).where(Reservation.id == res_id))
    r = res.scalar_one_or_none()
    if not r:
        raise HTTPException(status_code=404, detail="Reservation not found")
    for k, v in data.model_dump().items():
        setattr(r, k, v)
    await db.commit()
    return {"id": str(r.id), "status": "updated"}

@router.delete("/reservations/{res_id}", status_code=204)
async def delete_reservation(res_id: uuid.UUID, db: AsyncSession = Depends(get_db)):
    res = await db.execute(select(Reservation).where(Reservation.id == res_id))
    r = res.scalar_one_or_none()
    if r:
        await db.delete(r)
        await db.commit()
    return None


# -----------------------------------------------------------------------------
# Staff Task Board CRUD
# -----------------------------------------------------------------------------
@router.get("/tasks", response_model=List[LodgeTaskOut])
async def list_tasks(
    area: Optional[str] = None,
    complete: Optional[bool] = None,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    query = select(LodgeTask)
    if lodge_id:
        query = query.where(LodgeTask.lodge_id == lodge_id)
        
    query = query.order_by(LodgeTask.due_date.asc())
    if area and area != "all":
        query = query.where(LodgeTask.area == area)
    if complete is not None:
        query = query.where(LodgeTask.is_complete == complete)
    res = await db.execute(query)
    return res.scalars().all()

@router.post("/tasks", response_model=LodgeTaskOut, status_code=201)
async def create_task(
    data: LodgeTaskCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    t = LodgeTask(**data.model_dump())
    if lodge_id:
        t.lodge_id = lodge_id
    db.add(t)
    await db.commit()
    await db.refresh(t)
    return t

@router.put("/tasks/{task_id}", response_model=LodgeTaskOut)
async def update_task(
    task_id: uuid.UUID,
    data: LodgeTaskCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    stmt = select(LodgeTask).where(LodgeTask.id == task_id)
    if lodge_id:
        stmt = stmt.where(LodgeTask.lodge_id == lodge_id)
    res = await db.execute(stmt)
    t = res.scalar_one_or_none()
    if not t:
        raise HTTPException(status_code=404, detail="Task not found")
    
    was_complete = t.is_complete
    for k, v in data.model_dump().items():
        setattr(t, k, v)
    
    # If newly marked complete and has recurrence, auto-generate next task
    if not was_complete and t.is_complete and t.recurrence != "none" and t.due_date:
        next_date = t.due_date
        if t.recurrence == "daily":
            next_date += timedelta(days=1)
        elif t.recurrence == "weekly":
            next_date += timedelta(days=7)
        elif t.recurrence == "monthly":
            next_date += timedelta(days=30)
        
        next_task = LodgeTask(
            title=t.title,
            assigned_to=t.assigned_to,
            area=t.area,
            due_date=next_date,
            recurrence=t.recurrence,
            is_complete=False,
            notes=t.notes,
            lodge_id=lodge_id
        )
        db.add(next_task)
    
    await db.commit()
    await db.refresh(t)
    return t

@router.delete("/tasks/{task_id}", status_code=204)
async def delete_task(
    task_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    stmt = select(LodgeTask).where(LodgeTask.id == task_id)
    if lodge_id:
        stmt = stmt.where(LodgeTask.lodge_id == lodge_id)
    res = await db.execute(stmt)
    t = res.scalar_one_or_none()
    if t:
        await db.delete(t)
        await db.commit()
    return None


# -----------------------------------------------------------------------------
# Incidents CRUD
# -----------------------------------------------------------------------------
@router.get("/incidents", response_model=List[IncidentOut])
async def list_incidents(
    resolved: Optional[bool] = None,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    query = select(Incident)
    if lodge_id:
        query = query.where(Incident.lodge_id == lodge_id)
        
    query = query.order_by(
        desc(Incident.severity == "critical"),
        desc(Incident.created_at)
    )
    if resolved is not None:
        query = query.where(Incident.is_resolved == resolved)
    res = await db.execute(query)
    return res.scalars().all()

@router.post("/incidents", response_model=IncidentOut, status_code=201)
async def create_incident(
    data: IncidentCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    i = Incident(**data.model_dump())
    if lodge_id:
        i.lodge_id = lodge_id
    db.add(i)
    await db.commit()
    await db.refresh(i)
    return i

@router.put("/incidents/{inc_id}/resolve", response_model=IncidentOut)
async def resolve_incident(
    inc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    stmt = select(Incident).where(Incident.id == inc_id)
    if lodge_id:
        stmt = stmt.where(Incident.lodge_id == lodge_id)
    res = await db.execute(stmt)
    i = res.scalar_one_or_none()
    if not i:
        raise HTTPException(status_code=404, detail="Incident not found")
    i.is_resolved = True
    i.resolved_at = datetime.now()
    await db.commit()
    await db.refresh(i)
    return i

@router.delete("/incidents/{inc_id}", status_code=204)
async def delete_incident(
    inc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    stmt = select(Incident).where(Incident.id == inc_id)
    if lodge_id:
        stmt = stmt.where(Incident.lodge_id == lodge_id)
    res = await db.execute(stmt)
    i = res.scalar_one_or_none()
    if i:
        await db.delete(i)
        await db.commit()
    return None


# -----------------------------------------------------------------------------
# Daily Log CRUD
# -----------------------------------------------------------------------------
@router.get("/daily-log", response_model=List[DailyLogOut])
async def list_daily_logs(
    limit: int = 30,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    query = select(DailyLog)
    if lodge_id:
        query = query.where(DailyLog.lodge_id == lodge_id)
    res = await db.execute(query.order_by(desc(DailyLog.log_date)).limit(limit))
    return res.scalars().all()

@router.post("/daily-log", response_model=DailyLogOut, status_code=201)
async def create_or_update_daily_log(
    data: DailyLogCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    stmt = select(DailyLog).where(DailyLog.log_date == data.log_date)
    if lodge_id:
        stmt = stmt.where(DailyLog.lodge_id == lodge_id)
    res = await db.execute(stmt)
    existing = res.scalar_one_or_none()
    if existing:
        for k, v in data.model_dump().items():
            setattr(existing, k, v)
        await db.commit()
        await db.refresh(existing)
        return existing
    else:
        dl = DailyLog(**data.model_dump())
        if lodge_id:
            dl.lodge_id = lodge_id
        db.add(dl)
        await db.commit()
        await db.refresh(dl)
        return dl

# Rooms CRUD

@router.get("/rooms", response_model=List[RoomOut])
async def get_rooms(
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    stmt = select(Room)
    if lodge_id:
        stmt = stmt.where(Room.lodge_id == lodge_id)
    stmt = stmt.order_by(Room.created_at.asc())
    result = await db.execute(stmt)
    return result.scalars().all()

@router.post("/rooms", response_model=RoomOut, status_code=status.HTTP_201_CREATED)
async def create_room(
    data: RoomCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    if not lodge_id:
        raise HTTPException(status_code=400, detail="Lodge ID required")
    room = Room(name=data.name, lodge_id=lodge_id)
    db.add(room)
    await db.commit()
    await db.refresh(room)
    return room

@router.put("/rooms/{room_id}", response_model=RoomOut)
async def update_room(
    room_id: uuid.UUID,
    data: RoomCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    res = await db.execute(select(Room).where(Room.id == room_id, Room.lodge_id == lodge_id))
    room = res.scalar_one_or_none()
    if not room:
        raise HTTPException(status_code=404, detail="Room not found")
    room.name = data.name
    await db.commit()
    await db.refresh(room)
    return room

@router.delete("/rooms/{room_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_room(
    room_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    res = await db.execute(select(Room).where(Room.id == room_id, Room.lodge_id == lodge_id))
    room = res.scalar_one_or_none()
    if not room:
        raise HTTPException(status_code=404, detail="Room not found")
    await db.delete(room)
    await db.commit()
    return None

# Agencies CRUD

@router.get("/agencies", response_model=List[AgencyOut])
async def get_agencies(
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    stmt = select(Agency)
    if lodge_id:
        stmt = stmt.where(Agency.lodge_id == lodge_id)
    stmt = stmt.order_by(Agency.created_at.asc())
    result = await db.execute(stmt)
    return result.scalars().all()

@router.post("/agencies", response_model=AgencyOut, status_code=status.HTTP_201_CREATED)
async def create_agency(
    data: AgencyCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    if not lodge_id:
        raise HTTPException(status_code=400, detail="Lodge ID required")
    agency = Agency(name=data.name, color=data.color, lodge_id=lodge_id)
    db.add(agency)
    await db.commit()
    await db.refresh(agency)
    return agency

@router.put("/agencies/{agency_id}", response_model=AgencyOut)
async def update_agency(
    agency_id: uuid.UUID,
    data: AgencyCreate,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    res = await db.execute(select(Agency).where(Agency.id == agency_id, Agency.lodge_id == lodge_id))
    agency = res.scalar_one_or_none()
    if not agency:
        raise HTTPException(status_code=404, detail="Agency not found")
    agency.name = data.name
    agency.color = data.color
    await db.commit()
    await db.refresh(agency)
    return agency

@router.delete("/agencies/{agency_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_agency(
    agency_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    res = await db.execute(select(Agency).where(Agency.id == agency_id, Agency.lodge_id == lodge_id))
    agency = res.scalar_one_or_none()
    if not agency:
        raise HTTPException(status_code=404, detail="Agency not found")
    await db.delete(agency)
    await db.commit()
    return None

@router.get("/calendar/month-summary")
async def get_calendar_month_summary(
    month: str,
    user_email: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    import calendar
    from models import Recording, Reservation, Guest
    from sqlalchemy import or_, and_, select
    
    try:
        year, mth = map(int, month.split('-'))
        start_date = date(year, mth, 1)
        _, days_in_month = calendar.monthrange(year, mth)
        end_date = date(year, mth, days_in_month)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid month format")

    summary = {}
    for d in range(1, days_in_month + 1):
        d_str = date(year, mth, d).isoformat()
        summary[d_str] = {"bookings": 0, "memos": 0, "shopping": 0, "future_reminders": 0}

    # Fetch Bookings
    res_stmt = select(Reservation).join(Guest).where(
        Reservation.check_in <= end_date,
        Reservation.check_out >= start_date,
        Reservation.status != 'cancelled'
    )
    if lodge_id: res_stmt = res_stmt.where(Guest.lodge_id == lodge_id)
    reservations = (await db.execute(res_stmt)).scalars().all()
    
    for r in reservations:
        c_in = max(r.check_in, start_date)
        c_out = min(r.check_out or end_date, end_date)
        delta = (c_out - c_in).days
        for i in range(delta + 1):
            cur = c_in + timedelta(days=i)
            d_str = cur.isoformat()
            if d_str in summary:
                summary[d_str]["bookings"] += 1

    # Fetch Recordings
    rec_stmt = select(Recording).where(
        Recording.date_recorded >= start_date,
        Recording.date_recorded <= end_date
    )
    if lodge_id: rec_stmt = rec_stmt.where(Recording.lodge_id == lodge_id)
    if user_email: rec_stmt = rec_stmt.where(Recording.user_email == user_email)
    
    recordings = (await db.execute(rec_stmt)).scalars().all()
    
    for rec in recordings:
        if not rec.date_recorded: continue
        d_str = rec.date_recorded.isoformat()
        if d_str in summary:
            if rec.status == "done":
                if rec.type == "memo": summary[d_str]["memos"] += 1
                elif rec.type == "shopping": summary[d_str]["shopping"] += 1
            else:
                summary[d_str]["future_reminders"] += 1

    return summary


# Export Booking Sheet

@router.get("/export-booking-sheet")
async def export_booking_sheet(
    month: str, # YYYY-MM
    db: AsyncSession = Depends(get_db),
    lodge_id: Optional[uuid.UUID] = Depends(get_active_lodge_id)
):
    from fastapi.responses import StreamingResponse
    import io
    import calendar as pycal
    from openpyxl import Workbook
    from openpyxl.styles import PatternFill, Font, Alignment
    
    try:
        year, mth = map(int, month.split('-'))
        start_date = date(year, mth, 1)
        _, days_in_month = pycal.monthrange(year, mth)
        end_date = date(year, mth, days_in_month)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid month format")

    # Fetch Rooms
    room_stmt = select(Room)
    if lodge_id: room_stmt = room_stmt.where(Room.lodge_id == lodge_id)
    rooms = (await db.execute(room_stmt)).scalars().all()
    
    # Fetch Agencies
    agency_stmt = select(Agency)
    if lodge_id: agency_stmt = agency_stmt.where(Agency.lodge_id == lodge_id)
    agencies = (await db.execute(agency_stmt)).scalars().all()
    
    # Fetch Reservations
    res_stmt = select(Reservation).where(
        and_(
            Reservation.check_in <= end_date,
            Reservation.check_out >= start_date,
            Reservation.status != 'cancelled'
        )
    ).options(selectinload(Reservation.guest))
    reservations = (await db.execute(res_stmt)).scalars().all()
    
    wb = Workbook()
    ws = wb.active
    ws.title = f"Bookings {month}"
    
    # Legend Row (Row 2)
    ws.cell(row=2, column=1, value="Legend:")
    col_idx = 2
    for ag in agencies:
        c = ws.cell(row=2, column=col_idx, value=ag.name)
        if ag.color:
            hex_color = ag.color.replace('#', '')
            c.fill = PatternFill(start_color=hex_color, end_color=hex_color, fill_type="solid")
        col_idx += 1
        
    # Days Row (Row 4)
    ws.cell(row=4, column=1, value="Houses / Days").font = Font(bold=True)
    for day in range(1, days_in_month + 1):
        ws.cell(row=4, column=day + 1, value=day).font = Font(bold=True)
        
    # House Rows
    row_idx = 5
    for room in rooms:
        ws.cell(row=row_idx, column=1, value=room.name)
        # Check bookings for this room
        for day in range(1, days_in_month + 1):
            curr_date = date(year, mth, day)
            # Find a reservation covering this day and room
            match = None
            for r in reservations:
                # room names must match exactly or close enough. Here we use exact matching.
                if r.room_or_unit == room.name and r.check_in <= curr_date < r.check_out:
                    match = r
                    break
            
            if match:
                c = ws.cell(row=row_idx, column=day + 1, value=(match.guest.full_name if match.guest else 'Booked'))
                # Match agency color
                ag_color = None
                for ag in agencies:
                    if match.source and ag.name.lower() == match.source.lower() and ag.color:
                        ag_color = ag.color.replace('#', '')
                        break
                if ag_color:
                    c.fill = PatternFill(start_color=ag_color, end_color=ag_color, fill_type="solid")
                
        row_idx += 1
        
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    headers = {
        'Content-Disposition': f'attachment; filename="Booking_List_{month}.xlsx"'
    }
    return StreamingResponse(output, headers=headers, media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
